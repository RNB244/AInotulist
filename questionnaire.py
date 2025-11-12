from __future__ import annotations
from io import BytesIO
from typing import List, Dict, Tuple
from rapidfuzz import fuzz
from docx import Document
import re

def _normalize(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().lower())

def load_questions_from_docx(file_or_bytes) -> List[str]:
    """
    Leest een .docx vragenlijst. Pakt paragrafen die ofwel eindigen op '?' of beginnen met een nummering (1., 1), en filtert korte/lege regels.
    """
    if isinstance(file_or_bytes, (bytes, bytearray)):
        doc = Document(BytesIO(file_or_bytes))
    else:
        doc = Document(file_or_bytes)

    qs: List[str] = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        is_question = txt.endswith("?") or re.match(r"^\s*\d+[\).]\s+", txt)
        if is_question and len(txt) > 3:
            qs.append(txt)
    # fallback: als niets gevonden, neem alle niet-lege paragrafen
    if not qs:
        qs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return qs

def assign_segments_to_questions(
    segments: List[Dict],
    questions: List[str],
    threshold: int = 55,
    sequential: bool = True
) -> Tuple[Dict[int, List[Dict]], List[Tuple[Dict, int, int]]]:
    """
    Koppelt whisper-segmenten aan vragen m.b.v. fuzzy matching.
    - threshold: minimale matchscore (0-100)
    - sequential: vraagindex mag alleen gelijk blijven of oplopen
    Returns:
      mapping: {q_idx: [segment, ...]}
      debug_list: [(segment, best_q_idx, score)]
    """
    q_norm = [_normalize(q) for q in questions]
    mapping: Dict[int, List[Dict]] = {i: [] for i in range(len(questions))}
    debug: List[Tuple[Dict, int, int]] = []

    current_idx = 0
    for seg in segments:
        txt = _normalize(seg["text"])
        # kies alleen volgende 3 vragen als sequential aan staat
        candidate_range = range(current_idx, min(current_idx + 3, len(q_norm))) if sequential else range(len(q_norm))
        best_idx, best_score = -1, -1
        for i in candidate_range:
            score = fuzz.token_set_ratio(txt, q_norm[i])
            if score > best_score:
                best_idx, best_score = i, score

        # volgorde-constraint en drempel
        if best_idx >= 0 and best_score >= threshold:
            mapping[best_idx].append(seg)
            if sequential and best_idx >= current_idx:
                current_idx = best_idx
        else:
            # als geen goede match, hang aan huidige vraag (best effort)
            mapping[current_idx].append(seg)

        debug.append((seg, best_idx, best_score))
    return mapping, debug

def flatten_mapping_to_text(mapping: Dict[int, List[Dict]]) -> Dict[int, str]:
    """Combineert segments per vraag in één tekst."""
    combined: Dict[int, str] = {}
    for q_idx, segs in mapping.items():
        combined[q_idx] = " ".join(s["text"].strip() for s in segs if s.get("text"))
    return combined
