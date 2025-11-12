from __future__ import annotations
from io import BytesIO
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt

def build_docx_with_notes(
    title: str,
    questions: List[str],
    notes_per_question: Dict[int, List[str]],
    actions_per_question: Optional[Dict[int, List[str]]] = None,
    global_summary: str = "",
    global_actions: Optional[List[str]] = None
) -> bytes:
    """
    Genereert een .docx met alle vragen en notulen (bullets) per vraag.
    Actiepunten (indien aanwezig) komen onder een subkop per vraag.
    """
    doc = Document()
    doc.add_heading(title, level=1)

    # Per vraag: vraag + samenvatting als tekst
    for idx, q in enumerate(questions):
        doc.add_heading(f"Vraag {idx+1}", level=2)
        p = doc.add_paragraph(q)
        p.runs[0].font.size = Pt(11)
        # Samenvatting van gesprek (notulen als tekst)
        bullets = notes_per_question.get(idx, [])
        if bullets:
            doc.add_paragraph("Samenvatting van gesprek:", style=None)
            doc.add_paragraph(" ".join(bullets))
        else:
            doc.add_paragraph("Geen samenvatting beschikbaar.")

    # Aan het einde: alle actiepunten opgesomd
    all_actions = []
    if actions_per_question:
        for acts in actions_per_question.values():
            all_actions.extend(acts)
    if global_actions:
        all_actions.extend(global_actions)
    if all_actions:
        doc.add_heading("Actiepunten", level=2)
        for a in all_actions:
            doc.add_paragraph(a, style="List Bullet")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()
